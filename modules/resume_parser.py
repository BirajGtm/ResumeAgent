import os
import json
import re
from docx import Document
from modules.llm_router import run_llm_task
import docx2txt

INPUT_DIR = "input"
RAW_RESUME_PATH = os.path.join(INPUT_DIR, "uploaded_resume.txt")
PARSED_JSON_PATH = os.path.join(INPUT_DIR, "parsed_resume.json")
RESUME_DOCX_PATH = "cache/master_resume.docx"


def extract_text_with_docx2txt(docx_path):
    """
    Extracts all visible text from a DOCX file using docx2txt.
    This includes text boxes, headers, tables, etc., better than python-docx.
    """
    try:
        text = docx2txt.process(docx_path)
        return text.strip() if text else ""
    except Exception as e:
        return f"⚠️ Error extracting text: {e}"
    
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    # 1. Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # 2. Extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))  # Format as a readable line

    # 3. Extract header/footer text (if any)
    for section in doc.sections:
        header = section.header
        footer = section.footer
        if header and header.paragraphs:
            for para in header.paragraphs:
                if para.text.strip():
                    full_text.append("[Header] " + para.text.strip())
        if footer and footer.paragraphs:
            for para in footer.paragraphs:
                if para.text.strip():
                    full_text.append("[Footer] " + para.text.strip())

    return "\n".join(full_text)



def clean_llm_json_output(output):
    """
    Removes common LLM code block wrappers like ```json and ```
    """
    cleaned = re.sub(
        r"^```(?:json)?\s*|\s*```$", "", output.strip(), flags=re.IGNORECASE | re.MULTILINE
    )
    return cleaned.strip()

def get_parsed_resume(model, force_parse=False):
    """
    Loads cached parsed resume, or generates it via LLM if not found or forced.
    Also extracts resume text from .docx if needed.
    """
    # Decision: When should RAW_RESUME_PATH be updated from RESUME_DOCX_PATH?
    # - If RAW_RESUME_PATH doesn't exist AND RESUME_DOCX_PATH exists (initial run with DOCX).
    # - OR if force_parse is True AND RESUME_DOCX_PATH exists (new DOCX uploaded, so text needs refresh).
    # Ensure required directories exist
    os.makedirs("input", exist_ok=True)
    os.makedirs("cache", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    os.makedirs("temp_resume_data", exist_ok=True)

    should_extract_text_from_docx = (
        not os.path.exists(RAW_RESUME_PATH) and os.path.exists(RESUME_DOCX_PATH)
    ) or (force_parse and os.path.exists(RESUME_DOCX_PATH))

    if should_extract_text_from_docx:
        print(f"INFO: Extracting text from {RESUME_DOCX_PATH} to {RAW_RESUME_PATH} because should_extract_text_from_docx is True (force_parse: {force_parse})") # Logging for debug
        try:
            resume_text = extract_text_with_docx2txt(RESUME_DOCX_PATH)
            with open(RAW_RESUME_PATH, "w", encoding="utf-8") as f:
                f.write(resume_text)
        except Exception as e:
            return {"error": f"Failed to extract text from DOCX: {str(e)}"}


    if not os.path.exists(RAW_RESUME_PATH):
        return {"error": f"Resume text not found at {RAW_RESUME_PATH}. Ensure a resume is uploaded or exists."}

    # Now, handle cached JSON parsing or re-parsing
    if not force_parse and os.path.exists(PARSED_JSON_PATH):
        print(f"INFO: Attempting to load cached parsed resume from {PARSED_JSON_PATH}") # Logging for debug
        try:
            with open(PARSED_JSON_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            print(f"WARN: Failed to load cached JSON, falling back to re-parse. Error: {str(e)}") # Logging
            pass  # Fallback to re-parse

    # Force re-parse or parse for first time (if cache load failed or force_parse is true)
    print(f"INFO: Proceeding to parse resume text from {RAW_RESUME_PATH} with LLM (force_parse: {force_parse})") # Logging
    try:
        with open(RAW_RESUME_PATH, "r", encoding="utf-8") as f:
            resume_text = f.read()
        if not resume_text.strip():
            return {"error": f"Resume text file at {RAW_RESUME_PATH} is empty."}
        return parse_resume_with_llm(resume_text, model)
    except Exception as e:
        return {"error": f"Failed to read resume text for parsing: {str(e)}"}

def parse_resume_with_llm(resume_text, model):
    """
    Calls LLM to parse resume into structured JSON and saves it.
    """
    system_prompt = (
        "You are a highly precise resume sectioning and content extraction assistant. Your primary goal is to parse the provided resume text into a JSON object where the top-level keys are the EXACT section titles found in the input resume, and the values are the corresponding content from those sections.\n"
        "\n"
        "## Core Parsing Directives:\n"
        "1.  **Identify and Preserve Original Section Titles:** Your first task is to identify the main section headings used in the resume (e.g., 'Summary', 'Work Experience', 'Education', 'Skills', 'Projects', etc. - these will vary per resume). These exact titles, as they appear, will become the top-level keys in your JSON output.\n"
        "2.  **Extract Content Verbatim Under Each Section:** For each identified section, extract all the text content that falls under it. \n"
        "    -   If the content is a single block of text (like a 'Summary' or 'Objective'), the value for that section key should be a string.\n"
        "    -   If the content is a list of items (like bullet points under a 'Key Qualifications' section or a simple 'Skills' list), the value should be an array of strings, with each string being an item from the list. Preserve the original phrasing.\n"
        "    -   For complex sections like 'Work Experience', 'Projects', or 'Education' which typically contain multiple structured entries (e.g., multiple jobs, multiple projects, multiple degrees): \n"
        "        -   The value for such a section key should be an array of objects.\n"
        "        -   Each object within the array should represent one distinct entry (e.g., one job, one project, one degree).\n"
        "        -   Within each object, try to identify common sub-headings or distinct pieces of information (e.g., 'title', 'company', 'dates', 'responsibilities' for experience; 'project_name', 'description', 'technologies' for projects). Use these as keys within the sub-object. The 'responsibilities' or 'description' should often be an array of strings if they are bullet points in the original resume.\n"
        "3.  **Strict Adherence to Source - NO Modification or Inference:** Extract ONLY the information explicitly present. DO NOT rephrase, summarize, infer missing details, add information not present, or change the names of the sections found in the resume. Maintain the original structure of lists and bullet points within sections.\n"
        "4.  **Handle Contact Information (If Grouped):** If contact information (name, email, phone, LinkedIn, portfolio/website) is clearly grouped at the beginning of the resume, attempt to extract it into a top-level `contact_information` object. The individual's `name` should also be a top-level key if distinctly presented. If contact details are embedded within other sections or not clearly grouped, include them as found under their respective section, or the LLM should use its best judgment based on visual grouping if no explicit section title is present for contact info.\n"
        "\n"
        "❗**Output Format Constraint:** Return ONLY raw, valid JSON. Absolutely NO markdown formatting (e.g., ```json or ```).\n"
        "\n"
        "## Conceptual JSON Output Structure (This is illustrative; keys will depend on the specific resume's sections):\n"
        "```json\n"  # Note: This is an example for your understanding; it MUST NOT output the backticks.
        "{\n"
        '  "Full_Name": "Jane A. Doe",\n'
        '  "Contact_Details": {\n'
        '    "Mailing_Address": "123 Main St, Anytown, USA",\n'
        '    "Primary_Email": "jane.doe@email.com",\n'
        '    "Mobile_Number": "555-123-4567",\n'
        '    "LinkedIn_Profile": "linkedin.com/in/janedoe",\n'
        '    "Personal_Website": "janedoeportfolio.com"\n'
        '  },\n'
        '  "CAREER OBJECTIVE": "To secure a challenging position in a reputable organization...",\n'
        '  "PROFESSIONAL SKILLS": [\n'
        '    "Project Management",\n'
        '    "Data Analysis",\n'
        '    "Agile Methodologies"\n'
        '  ],\n'
        '  "EMPLOYMENT HISTORY": [\n'
        '    {\n'
        '      "Job_Title": "Senior Software Engineer",\n'
        '      "Company_Name": "Innovatech Solutions Inc.",\n'
        '      "Employment_Period": "June 2018 – Present",\n'
        '      "Key_Responsibilities": [\n'
        '        "Led a team of 5 developers in the design and implementation of new software features.",\n'
        '        "Reduced system downtime by 15% through proactive maintenance and monitoring."\n'
        '      ]\n'
        '    },\n'
        '    {\n'
        '      "Job_Title": "Software Developer",\n'
        '      "Company_Name": "Tech Systems Ltd.",\n'
        '      "Employment_Period": "July 2015 – May 2018",\n'
        '      "Key_Responsibilities": [\n'
        '        "Developed and maintained web applications using Java and Spring framework.",\n'
        '        "Contributed to all phases of the software development lifecycle."\n'
        '      ]\n'
        '    }\n'
        '  ],\n'
        '  "EDUCATION AND CERTIFICATIONS": [\n'  # Example if Education and Certs are grouped under one section title
        '    {\n'
        '      "Degree_Name": "Master of Science in Computer Science",\n'
        '      "Institution_Name": "State University",\n'
        '      "Graduation_Year": "2015",\n'
        '      "Relevant_Coursework": [\n'
        '         "Advanced Algorithms",\n'
        '         "Database Management Systems"\n'
        '       ]\n'
        '    },\n'
        '    {\n'
        '       "Certification_Name": "Certified Scrum Master (CSM)",\n'
        '       "Issuing_Body": "Scrum Alliance",\n'
        '       "Date_Obtained": "2019"\n'
        '    }\n'
        '  ]\n'
        '  // ... Any other sections identified in the resume will follow, using their original titles as keys.\n'
        "}\n"
        "```\n"  # End of example.
        "Your primary task is faithful sectioning based on the input resume's unique structure and then accurate content mirroring for each section. The JSON structure must adapt to the input resume's own organization and section titles."
    )

    user_prompt = (
        f"\nExtract structured data from this resume text and return it as clean JSON:\n\n---\n{resume_text}\n---\n"
    )

    response = run_llm_task(
        task="resume_structuring",
        prompt=user_prompt,
        context=None,
        model_preference=model,
        system_prompt=system_prompt,
    )

    raw_output = response.get("output", "{}")
    cleaned_output = clean_llm_json_output(raw_output)

    try:
        parsed = json.loads(cleaned_output)
        with open(PARSED_JSON_PATH, "w", encoding="utf-8") as f:
            json.dump(parsed, f, indent=2)
        return parsed
    except json.JSONDecodeError as e:
        return {
            "error": "⚠️ Failed to parse LLM output as JSON",
            "raw": raw_output,
            "cleaned": cleaned_output,
            "exception": str(e),
        }